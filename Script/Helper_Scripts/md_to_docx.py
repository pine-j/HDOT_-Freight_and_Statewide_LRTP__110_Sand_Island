"""
Generic Markdown to DOCX Converter

Converts any Markdown file to a formatted Word document.
Supports: headings, tables, bullet/numbered lists (nested), bold, italic, inline code,
          links, images, code blocks, blockquotes, horizontal rules.

Usage:
    python md_to_docx.py input.md [output.docx] [--author "Author Name"]
    
If output is not specified, it will use the same name as input with .docx extension.
If --author is not specified, the current OS username is used.
"""

import re
import sys
import os
import datetime
import traceback
from xml.sax.saxutils import escape as xml_escape
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
# When True, markdown horizontal rules (---, ***, ___) are rendered as a
# visible separator line in the Word document.  Set to False (default) to
# silently skip them.
ENABLE_HORIZONTAL_RULES = False


def add_hyperlink(paragraph, text, url):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )
    hyperlink = parse_xml(
        f'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'r:id="{r_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        f'<w:r><w:rPr><w:rStyle w:val="Hyperlink"/><w:color w:val="0563C1"/><w:u w:val="single"/></w:rPr>'
        f'<w:t>{xml_escape(text)}</w:t></w:r></w:hyperlink>'
    )
    paragraph._p.append(hyperlink)


def parse_inline_formatting(paragraph, text):
    """
    Parse inline markdown formatting and add to paragraph.
    Handles: `code`, **bold**, *italic*, [text](url)
    """
    if not text:
        return

    # Pattern matches (in priority order): inline code, links, bold, italic
    pattern = r'(`([^`]+)`|\[([^\]]+)\]\(([^)]+)\)|\*\*([^*]+)\*\*|\*([^*]+)\*)'

    last_end = 0
    for match in re.finditer(pattern, text):
        # Add any plain text before this match
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])

        full_match = match.group(0)

        if full_match.startswith('`'):
            # Inline code: `code`
            code_text = match.group(2)
            run = paragraph.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif full_match.startswith('['):
            # Link: [text](url)
            link_text = match.group(3)
            link_url = match.group(4)
            add_hyperlink(paragraph, link_text, link_url)
        elif full_match.startswith('**'):
            # Bold: **text**
            bold_text = match.group(5)
            run = paragraph.add_run(bold_text)
            run.bold = True
        elif full_match.startswith('*'):
            # Italic: *text*
            italic_text = match.group(6)
            run = paragraph.add_run(italic_text)
            run.italic = True

        last_end = match.end()

    # Add any remaining plain text after the last match
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def parse_table(lines, start_idx):
    """
    Parse a markdown table starting at the given index.
    Returns (table_data, alignments, end_index).
    - table_data: list of rows (each row is a list of cell strings)
    - alignments: list of alignment values ('left', 'center', 'right') per column
    """
    table_data = []
    alignments = []
    idx = start_idx

    while idx < len(lines):
        line = lines[idx].strip()

        if line.startswith('|') and line.endswith('|'):
            # Check if this is a separator row (contains only |, -, :, spaces)
            if re.match(r'^[\|\-:\s]+$', line):
                # Extract alignment from separator cells
                sep_cells = [cell.strip() for cell in line.split('|')[1:-1]]
                for cell in sep_cells:
                    if cell.startswith(':') and cell.endswith(':'):
                        alignments.append('center')
                    elif cell.endswith(':'):
                        alignments.append('right')
                    else:
                        alignments.append('left')
                idx += 1
                continue

            # Parse data cells
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            table_data.append(cells)
            idx += 1
        else:
            break

    return table_data, alignments, idx


def _get_alignment_enum(alignment_str):
    """Convert alignment string to WD_ALIGN_PARAGRAPH enum."""
    if alignment_str == 'center':
        return WD_ALIGN_PARAGRAPH.CENTER
    elif alignment_str == 'right':
        return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.LEFT


# ---------------------------------------------------------------------------
# Table layout helpers
# ---------------------------------------------------------------------------


def _get_available_page_width(doc):
    """Return the usable width between margins in inches."""
    try:
        section = doc.sections[0]
        return (section.page_width - section.left_margin - section.right_margin) / Inches(1)
    except (IndexError, TypeError):
        return 6.5  # Letter with 1-inch margins


# Default cell padding in twips (1 inch = 1440 twips, 1 pt = 20 twips).
_CELL_MARGIN_TOP_TWIPS = 40      # ~2 pt / 0.028 in
_CELL_MARGIN_BOTTOM_TWIPS = 40   # ~2 pt / 0.028 in


def _set_cell_margins(cell, top=_CELL_MARGIN_TOP_TWIPS,
                      bottom=_CELL_MARGIN_BOTTOM_TWIPS):
    """Add top/bottom padding to a DOCX table cell via XML.

    Values are in twips (twentieths of a point).  The defaults give a
    comfortable ~4 pt cushion above and below the text.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # Remove any existing <w:tcMar> so we don't duplicate
    for existing in tcPr.findall(f'{{{ns}}}tcMar'):
        tcPr.remove(existing)

    tcMar = parse_xml(
        f'<w:tcMar xmlns:w="{ns}">'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


# Constants for column width estimation
_CHAR_WIDTH_INCHES = 0.08       # Approx. width per char at 11 pt Calibri
_CELL_HORIZ_PAD_INCHES = 0.25   # Horizontal padding inside a cell
_BOLD_WIDTH_FACTOR = 1.2        # Bold text is ~20 % wider
_EXTRA_SPACE_POWER = 1.35       # >1 biases extra width to verbose columns
_HEADER_WIDTH_BOOSTS = {
    'period': 1.65,
    'description': 1.50,
    'details': 1.35,
    'notes': 1.30,
    'criteria': 1.25,
    'rule': 1.20,
    'logic': 1.20,
    'impact': 1.20,
    'mitigation': 1.20,
    'analysis': 1.20,
}


def _strip_to_rendered_text(cell_text):
    """Strip Markdown syntax that is not rendered, keeping all visible chars.

    Unlike the scoring helper (which also strips brackets/parens), this
    preserves parentheses and brackets that appear literally in the cell
    because they occupy space in the rendered Word document.
    """
    # Remove bold/italic markers and backtick code spans
    text = re.sub(r'[*`]', '', cell_text)
    # Convert link syntax to visible text: [text](url) -> text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\\1', text)
    return text


def _min_column_widths_for_words(table_data, num_cols):
    """Compute minimum column widths (inches) so that the longest word
    in each column can render on a single line without mid-word breaks.

    "Word" here means a whitespace-delimited token — the smallest unit
    that Word will not break internally (hyphens and slashes may still
    act as break opportunities, but we are conservative).

    Returns:
        List of minimum widths in inches, one per column.
    """
    min_widths = []
    for col_idx in range(num_cols):
        max_word_width = 0.0
        for row_idx, row in enumerate(table_data):
            if col_idx < len(row):
                raw = row[col_idx]
                text = _strip_to_rendered_text(raw)
                # Detect if this cell renders bold (header row or **…**)
                is_bold = (row_idx == 0) or ('**' in raw)
                for word in text.split():
                    w = len(word) * _CHAR_WIDTH_INCHES
                    if is_bold:
                        w *= _BOLD_WIDTH_FACTOR
                    max_word_width = max(max_word_width, w)
        min_widths.append(max_word_width + _CELL_HORIZ_PAD_INCHES)
    return min_widths


def _natural_column_widths(table_data, num_cols):
    """Compute the natural width (inches) each column needs to display its
    longest cell content on a single line without any wrapping.

    This gives the "ideal" width — the column is wide enough that no cell
    text wraps.  The value is used as the upper target when allocating space.
    """
    natural = []
    for col_idx in range(num_cols):
        max_width = 0.0
        for row_idx, row in enumerate(table_data):
            if col_idx < len(row):
                text = _strip_to_rendered_text(row[col_idx])
                is_bold = (row_idx == 0) or ('**' in row[col_idx])
                w = len(text) * _CHAR_WIDTH_INCHES
                if is_bold:
                    w *= _BOLD_WIDTH_FACTOR
                max_width = max(max_width, w)
        natural.append(max_width + _CELL_HORIZ_PAD_INCHES)
    return natural


def _column_extra_space_weights(table_data, natural, minimums, num_cols):
    """Weight extra space allocation toward columns likely to wrap heavily."""
    weights = []
    header_row = table_data[0] if table_data else []

    for col_idx in range(num_cols):
        n = natural[col_idx]
        m = minimums[col_idx]
        desire = max(0.0, n - m)
        if desire <= 0:
            weights.append(0.0)
            continue

        # Prefer columns far from their minimum and with higher wrap pressure.
        wrap_pressure = (n / m) if m > 0 else 1.0
        weight = (desire ** _EXTRA_SPACE_POWER) * wrap_pressure

        # Common long-text headers deserve a little more width.
        header_text = ''
        if col_idx < len(header_row):
            header_text = _strip_to_rendered_text(header_row[col_idx]).lower()
        boost = 1.0
        for key, factor in _HEADER_WIDTH_BOOSTS.items():
            if key in header_text:
                boost = max(boost, factor)
        weight *= boost

        weights.append(weight)

    return weights


def _compute_column_widths(table_data, num_cols, available_width_inches):
    """Compute absolute column widths (inches) using a natural-width-based
    allocation strategy.

    Strategy
    --------
    1. Compute the *natural* width each column needs (longest cell on one
       line) and the *minimum* width (longest single word).
    2. If all columns fit at natural width, use those widths and distribute
       any surplus proportionally — the table looks great with no wrapping.
    3. If some wrapping is required, start every column at its minimum width
       and distribute the remaining space proportionally to each column's
       *desire* for more room (natural − minimum).  Columns with longer
       content receive more of the extra space; short-label columns stay
       compact.
    4. The result is normalised so the table spans exactly the available
       width.
    """
    natural = _natural_column_widths(table_data, num_cols)
    minimums = _min_column_widths_for_words(table_data, num_cols)

    total_natural = sum(natural)
    total_min = sum(minimums)

    # --- Case 1: every column fits on one line --------------------------
    if total_natural <= available_width_inches:
        widths = natural[:]
        surplus = available_width_inches - total_natural
        if total_natural > 0:
            widths = [w + surplus * (w / total_natural) for w in widths]
        return widths

    # --- Case 2: even single-word minimums exceed the page --------------
    if total_min >= available_width_inches:
        factor = (available_width_inches / total_min) if total_min > 0 else 1.0
        return [m * factor for m in minimums]

    # --- Case 3: normal — allocate extra space above minimums -----------
    extra = available_width_inches - total_min
    weights = _column_extra_space_weights(
        table_data, natural, minimums, num_cols)
    total_weight = sum(weights)

    if total_weight > 0:
        widths = [m + extra * (w / total_weight)
                  for m, w in zip(minimums, weights)]
    else:
        widths = [m + extra / num_cols for m in minimums]

    return widths


def _prevent_row_splitting(table):
    """Prevent table rows from splitting across page boundaries."""
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        if trPr.find(f'{{{ns}}}cantSplit') is None:
            trPr.append(parse_xml(f'<w:cantSplit xmlns:w="{ns}"/>'))


def _apply_fixed_table_layout(table, table_data, num_cols,
                              available_width_inches):
    """Apply a fixed-width, full-page-width table layout.

    * The table spans the full width between page margins.
    * Column widths are computed using a natural-width-based strategy:
      short-label columns stay compact while longer-content columns
      receive more space.
    * Every column is guaranteed wide enough for its longest single
      word to render without a mid-word line break.
    * Word is told to respect the explicit column widths (fixed layout).
    """
    widths = _compute_column_widths(table_data, num_cols, available_width_inches)

    # Normalize to exactly fill the available width
    total = sum(widths)
    if total > 0:
        factor = available_width_inches / total
        widths = [w * factor for w in widths]

    # ------------------------------------------------------------------
    # Apply XML properties
    # ------------------------------------------------------------------
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    tblPr = table._tbl.tblPr

    # Table width -> 100 % of page (pct 5000 = 100 %)
    for child in tblPr.findall(f'{{{ns}}}tblW'):
        tblPr.remove(child)
    tblPr.append(parse_xml(
        f'<w:tblW xmlns:w="{ns}" w:type="pct" w:w="5000"/>'
    ))

    # Fixed layout -> Word respects explicit column widths
    for child in tblPr.findall(f'{{{ns}}}tblLayout'):
        tblPr.remove(child)
    tblPr.append(parse_xml(
        f'<w:tblLayout xmlns:w="{ns}" w:type="fixed"/>'
    ))

    # Disable python-docx autofit flag
    table.autofit = False

    # Avoid ugly split rows at page boundaries.
    _prevent_row_splitting(table)

    # Set explicit column (and cell) widths
    for col_idx in range(num_cols):
        col_width = Inches(widths[col_idx])
        table.columns[col_idx].width = col_width
        for row in table.rows:
            row.cells[col_idx].width = col_width


def _is_special_line(stripped):
    """Check if a stripped line is a special Markdown element (not a regular paragraph line)."""
    if not stripped:
        return True
    if re.match(r'^(-{3,}|\*{3,}|_{3,})$', stripped):
        return True  # Horizontal rule
    if re.match(r'^#{1,6}\s+', stripped):
        return True  # Heading
    if stripped.startswith('|') and '|' in stripped[1:]:
        return True  # Table
    if re.match(r'^[-*+]\s+', stripped):
        return True  # Bullet list
    if re.match(r'^\d+[.)]\s+', stripped):
        return True  # Numbered list
    if stripped.startswith('```'):
        return True  # Code fence
    if stripped.startswith('>'):
        return True  # Blockquote
    if re.match(r'^!\[([^\]]*)\]\(([^)]+)\)', stripped):
        return True  # Image
    return False


def _restart_list_numbering(doc, paragraph):
    """Restart numbering at 1 for the first paragraph of a new numbered list.

    Creates a new numbering instance (w:num) in the document's numbering
    definitions that references the same abstract definition as the list
    style but includes a level-override to restart at value 1.  The
    paragraph's numPr is then pointed at the new instance.
    """
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # --- Locate the numId that the paragraph's list style uses ----------
    num_id = None
    ilvl = 0

    # Try paragraph-level numPr first
    pPr = paragraph._p.pPr
    if pPr is not None:
        numPr = pPr.find(f'{{{ns}}}numPr')
        if numPr is not None:
            el = numPr.find(f'{{{ns}}}numId')
            if el is not None:
                num_id = int(el.get(f'{{{ns}}}val'))
            el = numPr.find(f'{{{ns}}}ilvl')
            if el is not None:
                ilvl = int(el.get(f'{{{ns}}}val'))

    # Fallback: check the style element
    if num_id is None:
        style_elm = paragraph.style.element
        style_pPr = style_elm.find(f'{{{ns}}}pPr')
        if style_pPr is not None:
            numPr = style_pPr.find(f'{{{ns}}}numPr')
            if numPr is not None:
                el = numPr.find(f'{{{ns}}}numId')
                if el is not None:
                    num_id = int(el.get(f'{{{ns}}}val'))
                el = numPr.find(f'{{{ns}}}ilvl')
                if el is not None:
                    ilvl = int(el.get(f'{{{ns}}}val'))

    if num_id is None:
        return  # Cannot determine numbering

    # --- Find the abstractNumId for this numId --------------------------
    numbering_elm = doc.part.numbering_part._element
    abstract_num_id = None
    for num in numbering_elm.findall(f'{{{ns}}}num'):
        if int(num.get(f'{{{ns}}}numId')) == num_id:
            anid = num.find(f'{{{ns}}}abstractNumId')
            if anid is not None:
                abstract_num_id = int(anid.get(f'{{{ns}}}val'))
            break

    if abstract_num_id is None:
        return

    # --- Create a new w:num with a startOverride ------------------------
    max_num_id = max(
        (int(n.get(f'{{{ns}}}numId'))
         for n in numbering_elm.findall(f'{{{ns}}}num')),
        default=0,
    )
    new_num_id = max_num_id + 1

    new_num = parse_xml(
        f'<w:num xmlns:w="{ns}" w:numId="{new_num_id}">'
        f'<w:abstractNumId w:val="{abstract_num_id}"/>'
        f'<w:lvlOverride w:ilvl="{ilvl}">'
        f'<w:startOverride w:val="1"/>'
        f'</w:lvlOverride>'
        f'</w:num>'
    )
    numbering_elm.append(new_num)

    # --- Point the paragraph at the new numId ---------------------------
    pPr = paragraph._p.get_or_add_pPr()
    for existing in pPr.findall(f'{{{ns}}}numPr'):
        pPr.remove(existing)

    new_numPr = parse_xml(
        f'<w:numPr xmlns:w="{ns}">'
        f'<w:ilvl w:val="{ilvl}"/>'
        f'<w:numId w:val="{new_num_id}"/>'
        f'</w:numPr>'
    )
    pPr.insert(0, new_numPr)


def _add_code_block(doc, code_lines):
    """Add a code block to the document with monospace formatting and shading."""
    code_text = '\n'.join(code_lines)
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    # Add a light gray background shading to the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(
        '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'w:fill="F2F2F2" w:val="clear"/>'
    )
    pPr.append(shd)

    # Add indent for visual separation
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)


def _clean_document_metadata(doc, author):
    """Remove all python-docx traces from document metadata.

    Sets the author, clears comments, and updates timestamps so the
    document looks like it was created natively in Microsoft Word.
    """
    # --- Core properties (visible in File > Info) ----------------------
    props = doc.core_properties
    props.author = author
    props.last_modified_by = author
    props.comments = ''
    props.title = ''
    props.subject = ''
    props.keywords = ''
    props.category = ''
    props.created = datetime.datetime.now()
    props.modified = datetime.datetime.now()
    props.revision = 1

    # --- Extended properties (app.xml) ---------------------------------
    # Replace "python-docx" application identifier with MS Word values.
    try:
        from lxml import etree as _etree
        ep_reltype = ('http://schemas.openxmlformats.org/officeDocument/'
                      '2006/relationships/extended-properties')
        for rel in doc.part.package.rels.values():
            if rel.reltype == ep_reltype:
                part = rel.target_part
                root = _etree.fromstring(part.blob)
                ep_ns = ('http://schemas.openxmlformats.org/officeDocument/'
                         '2006/extended-properties')
                app_elem = root.find(f'{{{ep_ns}}}Application')
                if app_elem is not None:
                    app_elem.text = 'Microsoft Office Word'
                ver_elem = root.find(f'{{{ep_ns}}}AppVersion')
                if ver_elem is not None:
                    ver_elem.text = '16.0000'
                part._blob = _etree.tostring(root, xml_declaration=True,
                                             encoding='UTF-8',
                                             standalone=True)
                break
    except Exception:
        pass  # Non-critical; core properties cover the visible fields


def convert_markdown_to_docx(input_path, output_path=None, author=None):
    """
    Convert a markdown file to a Word document.
    
    Args:
        input_path: Path to the input markdown file
        output_path: Path for the output docx file (optional)
        author: Author name for document properties (defaults to OS username)
    """
    if output_path is None:
        output_path = os.path.splitext(input_path)[0] + '.docx'

    # Read the markdown file
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    doc = Document()

    # Clean metadata — remove python-docx traces, set author
    if author is None:
        author = os.environ.get('USERNAME',
                                os.environ.get('USER', ''))
    _clean_document_metadata(doc, author)

    idx = 0
    has_content = False  # Track whether any content has been added (for page-break logic)
    first_chapter_done = False  # Allow the first H2 to stay on the same page as the title
    prev_was_numbered_list = False  # Track numbered-list continuity

    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        # Skip empty lines (don't reset list tracking)
        if not stripped:
            idx += 1
            continue

        # Save and reset numbered-list tracking for this iteration.
        # The numbered-list branch below will set it back to True.
        was_in_numbered_list = prev_was_numbered_list
        prev_was_numbered_list = False

        # --- Fenced code block ---
        code_fence_match = re.match(r'^```(\w*)$', stripped)
        if code_fence_match:
            code_lines = []
            idx += 1
            while idx < len(lines):
                if lines[idx].strip() == '```':
                    idx += 1
                    break
                code_lines.append(lines[idx])
                idx += 1
            _add_code_block(doc, code_lines)
            has_content = True
            continue

        # --- Horizontal rule (must use 3+ of the SAME character) ---
        if re.match(r'^(-{3,}|\*{3,}|_{3,})$', stripped):
            if ENABLE_HORIZONTAL_RULES:
                p = doc.add_paragraph()
                p.add_run('─' * 50)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                has_content = True
            idx += 1
            continue

        # --- Headings ---
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            # Remove any trailing # characters
            heading_text = re.sub(r'\s*#+\s*$', '', heading_text)
            # Start each chapter (H1 or H2) on a new page.
            # Uses an explicit page-break paragraph (not the paragraph
            # property page_break_before) so the user can select and
            # delete it in Word to merge chapters.
            # - Skip when no content yet (avoids a leading blank page).
            # - Let the first H2 share the page with the title.
            if level <= 2 and has_content:
                if level == 1 or first_chapter_done:
                    doc.add_page_break()
            heading = doc.add_heading(heading_text, level=level)
            if level == 1:
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if level == 2:
                first_chapter_done = True
            has_content = True
            idx += 1
            continue

        # --- Image (block-level) ---
        image_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)$', stripped)
        if image_match:
            alt_text = image_match.group(1)
            image_path = image_match.group(2)
            # Resolve relative paths against the input file's directory
            if not os.path.isabs(image_path):
                input_dir = os.path.dirname(os.path.abspath(input_path))
                image_path = os.path.join(input_dir, image_path)
            try:
                doc.add_picture(image_path, width=Inches(5.5))
                # Center the image
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                # If image can't be loaded, add alt text as fallback
                p = doc.add_paragraph()
                run = p.add_run(f'[Image: {alt_text}]')
                run.italic = True
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            has_content = True
            idx += 1
            continue

        # --- Table ---
        if stripped.startswith('|') and '|' in stripped[1:]:
            table_data, alignments, new_idx = parse_table(lines, idx)

            if table_data:
                idx = new_idx
                # Create table
                num_cols = max(len(row) for row in table_data)
                table = doc.add_table(rows=len(table_data), cols=num_cols)
                table.style = 'Table Grid'

                # Pad alignments list to match column count
                while len(alignments) < num_cols:
                    alignments.append('left')

                available_width = _get_available_page_width(doc)

                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < num_cols:
                            cell = table.rows[row_idx].cells[col_idx]
                            # Clear existing paragraph and add formatted text
                            cell.text = ''
                            p = cell.paragraphs[0]
                            parse_inline_formatting(p, cell_text)

                            # Apply column alignment
                            p.alignment = _get_alignment_enum(alignments[col_idx])

                            # Vertical center alignment and cell padding
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            _set_cell_margins(cell)

                            # Bold header row
                            if row_idx == 0:
                                for run in p.runs:
                                    run.bold = True

                # Apply fixed layout with content-proportional column widths
                _apply_fixed_table_layout(
                    table, table_data, num_cols, available_width)

                doc.add_paragraph()  # Add spacing after table
                has_content = True
                continue

        # --- Blockquote ---
        if stripped.startswith('>'):
            quote_lines = []
            while idx < len(lines):
                s = lines[idx].strip()
                if s.startswith('>'):
                    # Remove the > prefix and optional following space
                    quote_text = re.sub(r'^>\s?', '', s)
                    quote_lines.append(quote_text)
                    idx += 1
                elif s == '':
                    # Allow blank lines within a blockquote if next line continues it
                    if idx + 1 < len(lines) and lines[idx + 1].strip().startswith('>'):
                        quote_lines.append('')
                        idx += 1
                    else:
                        break
                else:
                    break

            # Merge non-empty lines into a single paragraph
            merged_quote = ' '.join(ln for ln in quote_lines if ln)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            parse_inline_formatting(p, merged_quote)

            # Style the blockquote runs (italic + gray)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            # Add a left border via XML
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:left w:val="single" w:sz="12" w:space="4" w:color="AAAAAA"/>'
                '</w:pBdr>'
            )
            pPr.append(pBdr)
            has_content = True
            continue

        # --- Bullet list (with nesting support) ---
        bullet_match = re.match(r'^(\s*)([-*+])\s+(.+)$', line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            # Determine nesting level based on indentation
            if indent >= 4:
                style = 'List Bullet 3'
            elif indent >= 2:
                style = 'List Bullet 2'
            else:
                style = 'List Bullet'
            p = doc.add_paragraph(style=style)
            parse_inline_formatting(p, bullet_match.group(3))
            has_content = True
            idx += 1
            continue

        # --- Numbered list (with nesting support) ---
        number_match = re.match(r'^(\s*)(\d+)[.)]\s+(.+)$', line)
        if number_match:
            indent = len(number_match.group(1))
            if indent >= 4:
                style = 'List Number 3'
            elif indent >= 2:
                style = 'List Number 2'
            else:
                style = 'List Number'
            p = doc.add_paragraph(style=style)
            parse_inline_formatting(p, number_match.group(3))
            # Restart numbering at 1 if this is the first item of a new list
            if not was_in_numbered_list:
                _restart_list_numbering(doc, p)
            prev_was_numbered_list = True
            has_content = True
            idx += 1
            continue

        # --- Regular paragraph (merge consecutive non-special lines) ---
        para_lines = [stripped]
        idx += 1
        while idx < len(lines):
            next_stripped = lines[idx].strip()
            if not next_stripped or _is_special_line(next_stripped):
                break
            para_lines.append(next_stripped)
            idx += 1

        merged_text = ' '.join(para_lines)
        p = doc.add_paragraph()
        parse_inline_formatting(p, merged_text)
        has_content = True

    # Save the document
    doc.save(output_path)
    print(f"Successfully converted: {input_path}")
    print(f"Output saved to: {output_path}")
    return output_path


def main():
    """Main entry point for command-line usage."""
    # Parse optional --author flag
    author = None
    args = list(sys.argv[1:])
    if '--author' in args:
        ai = args.index('--author')
        if ai + 1 < len(args):
            author = args[ai + 1]
            del args[ai:ai + 2]
        else:
            print("Error: --author requires a name argument.")
            sys.exit(1)

    if len(args) < 1:
        print(__doc__)
        print("\nError: Please provide an input markdown file.")
        print("Usage: python md_to_docx.py input.md [output.docx] "
              "[--author \"Author Name\"]")
        sys.exit(1)

    input_file = args[0]
    output_file = args[1] if len(args) > 1 else None

    if not os.path.exists(input_file):
        print(f"Error: Input file '{input_file}' not found.")
        sys.exit(1)

    try:
        convert_markdown_to_docx(input_file, output_file, author=author)
    except Exception as e:
        print(f"Error during conversion: {e}")
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
