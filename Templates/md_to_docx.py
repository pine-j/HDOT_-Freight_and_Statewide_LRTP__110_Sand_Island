"""
Generic Markdown to DOCX Converter

Converts any Markdown file to a formatted Word document.
Supports: headings, tables, bullet/numbered lists (nested), bold, italic, inline code,
          links, images, code blocks, blockquotes, horizontal rules.

Usage:
    python md_to_docx.py input.md [output.docx]
    
If output is not specified, it will use the same name as input with .docx extension.
"""

import re
import sys
import os
import traceback
from xml.sax.saxutils import escape as xml_escape
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml


def add_hyperlink(paragraph, text, url):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )
    hyperlink = parse_xml(
        f'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'r:id="{r_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        f'<w:r><w:rPr><w:rStyle w:val="Hyperlink"/><w:color w:val="0563C1"/><w:u w:val="single"/></w:rPr>'
        f'<w:t>{xml_escape(text)}</w:t></w:r></w:hyperlink>'
    )
    paragraph._p.append(hyperlink)


def parse_inline_formatting(paragraph, text):
    """
    Parse inline markdown formatting and add to paragraph.
    Handles: `code`, **bold**, *italic*, [text](url)
    """
    if not text:
        return

    # Pattern matches (in priority order): inline code, links, bold, italic
    pattern = r'(`([^`]+)`|\[([^\]]+)\]\(([^)]+)\)|\*\*([^*]+)\*\*|\*([^*]+)\*)'

    last_end = 0
    for match in re.finditer(pattern, text):
        # Add any plain text before this match
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])

        full_match = match.group(0)

        if full_match.startswith('`'):
            # Inline code: `code`
            code_text = match.group(2)
            run = paragraph.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif full_match.startswith('['):
            # Link: [text](url)
            link_text = match.group(3)
            link_url = match.group(4)
            add_hyperlink(paragraph, link_text, link_url)
        elif full_match.startswith('**'):
            # Bold: **text**
            bold_text = match.group(5)
            run = paragraph.add_run(bold_text)
            run.bold = True
        elif full_match.startswith('*'):
            # Italic: *text*
            italic_text = match.group(6)
            run = paragraph.add_run(italic_text)
            run.italic = True

        last_end = match.end()

    # Add any remaining plain text after the last match
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def parse_table(lines, start_idx):
    """
    Parse a markdown table starting at the given index.
    Returns (table_data, alignments, end_index).
    - table_data: list of rows (each row is a list of cell strings)
    - alignments: list of alignment values ('left', 'center', 'right') per column
    """
    table_data = []
    alignments = []
    idx = start_idx

    while idx < len(lines):
        line = lines[idx].strip()

        if line.startswith('|') and line.endswith('|'):
            # Check if this is a separator row (contains only |, -, :, spaces)
            if re.match(r'^[\|\-:\s]+$', line):
                # Extract alignment from separator cells
                sep_cells = [cell.strip() for cell in line.split('|')[1:-1]]
                for cell in sep_cells:
                    if cell.startswith(':') and cell.endswith(':'):
                        alignments.append('center')
                    elif cell.endswith(':'):
                        alignments.append('right')
                    else:
                        alignments.append('left')
                idx += 1
                continue

            # Parse data cells
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            table_data.append(cells)
            idx += 1
        else:
            break

    return table_data, alignments, idx


def _get_alignment_enum(alignment_str):
    """Convert alignment string to WD_ALIGN_PARAGRAPH enum."""
    if alignment_str == 'center':
        return WD_ALIGN_PARAGRAPH.CENTER
    elif alignment_str == 'right':
        return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.LEFT


def _is_special_line(stripped):
    """Check if a stripped line is a special Markdown element (not a regular paragraph line)."""
    if not stripped:
        return True
    if re.match(r'^(-{3,}|\*{3,}|_{3,})$', stripped):
        return True  # Horizontal rule
    if re.match(r'^#{1,6}\s+', stripped):
        return True  # Heading
    if stripped.startswith('|') and '|' in stripped[1:]:
        return True  # Table
    if re.match(r'^[-*+]\s+', stripped):
        return True  # Bullet list
    if re.match(r'^\d+[.)]\s+', stripped):
        return True  # Numbered list
    if stripped.startswith('```'):
        return True  # Code fence
    if stripped.startswith('>'):
        return True  # Blockquote
    if re.match(r'^!\[([^\]]*)\]\(([^)]+)\)', stripped):
        return True  # Image
    return False


def _add_code_block(doc, code_lines):
    """Add a code block to the document with monospace formatting and shading."""
    code_text = '\n'.join(code_lines)
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    # Add a light gray background shading to the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(
        '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'w:fill="F2F2F2" w:val="clear"/>'
    )
    pPr.append(shd)

    # Add indent for visual separation
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)


def convert_markdown_to_docx(input_path, output_path=None):
    """
    Convert a markdown file to a Word document.
    
    Args:
        input_path: Path to the input markdown file
        output_path: Path for the output docx file (optional)
    """
    if output_path is None:
        output_path = os.path.splitext(input_path)[0] + '.docx'

    # Read the markdown file
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    doc = Document()

    idx = 0

    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            idx += 1
            continue

        # --- Fenced code block ---
        code_fence_match = re.match(r'^```(\w*)$', stripped)
        if code_fence_match:
            code_lines = []
            idx += 1
            while idx < len(lines):
                if lines[idx].strip() == '```':
                    idx += 1
                    break
                code_lines.append(lines[idx])
                idx += 1
            _add_code_block(doc, code_lines)
            continue

        # --- Horizontal rule (must use 3+ of the SAME character) ---
        if re.match(r'^(-{3,}|\*{3,}|_{3,})$', stripped):
            p = doc.add_paragraph()
            p.add_run('─' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            idx += 1
            continue

        # --- Headings ---
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            # Remove any trailing # characters
            heading_text = re.sub(r'\s*#+\s*$', '', heading_text)
            heading = doc.add_heading(heading_text, level=level)
            if level == 1:
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            idx += 1
            continue

        # --- Image (block-level) ---
        image_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)$', stripped)
        if image_match:
            alt_text = image_match.group(1)
            image_path = image_match.group(2)
            # Resolve relative paths against the input file's directory
            if not os.path.isabs(image_path):
                input_dir = os.path.dirname(os.path.abspath(input_path))
                image_path = os.path.join(input_dir, image_path)
            try:
                doc.add_picture(image_path, width=Inches(5.5))
                # Center the image
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                # If image can't be loaded, add alt text as fallback
                p = doc.add_paragraph()
                run = p.add_run(f'[Image: {alt_text}]')
                run.italic = True
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            idx += 1
            continue

        # --- Table ---
        if stripped.startswith('|') and '|' in stripped[1:]:
            table_data, alignments, new_idx = parse_table(lines, idx)

            if table_data:
                idx = new_idx
                # Create table
                num_cols = max(len(row) for row in table_data)
                table = doc.add_table(rows=len(table_data), cols=num_cols)
                table.style = 'Table Grid'

                # Pad alignments list to match column count
                while len(alignments) < num_cols:
                    alignments.append('left')

                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < num_cols:
                            cell = table.rows[row_idx].cells[col_idx]
                            # Clear existing paragraph and add formatted text
                            cell.text = ''
                            p = cell.paragraphs[0]
                            parse_inline_formatting(p, cell_text)

                            # Apply column alignment
                            p.alignment = _get_alignment_enum(alignments[col_idx])

                            # Bold header row
                            if row_idx == 0:
                                for run in p.runs:
                                    run.bold = True

                doc.add_paragraph()  # Add spacing after table
                continue

        # --- Blockquote ---
        if stripped.startswith('>'):
            quote_lines = []
            while idx < len(lines):
                s = lines[idx].strip()
                if s.startswith('>'):
                    # Remove the > prefix and optional following space
                    quote_text = re.sub(r'^>\s?', '', s)
                    quote_lines.append(quote_text)
                    idx += 1
                elif s == '':
                    # Allow blank lines within a blockquote if next line continues it
                    if idx + 1 < len(lines) and lines[idx + 1].strip().startswith('>'):
                        quote_lines.append('')
                        idx += 1
                    else:
                        break
                else:
                    break

            # Merge non-empty lines into a single paragraph
            merged_quote = ' '.join(ln for ln in quote_lines if ln)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            parse_inline_formatting(p, merged_quote)

            # Style the blockquote runs (italic + gray)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            # Add a left border via XML
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:left w:val="single" w:sz="12" w:space="4" w:color="AAAAAA"/>'
                '</w:pBdr>'
            )
            pPr.append(pBdr)
            continue

        # --- Bullet list (with nesting support) ---
        bullet_match = re.match(r'^(\s*)([-*+])\s+(.+)$', line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            # Determine nesting level based on indentation
            if indent >= 4:
                style = 'List Bullet 3'
            elif indent >= 2:
                style = 'List Bullet 2'
            else:
                style = 'List Bullet'
            p = doc.add_paragraph(style=style)
            parse_inline_formatting(p, bullet_match.group(3))
            idx += 1
            continue

        # --- Numbered list (with nesting support) ---
        number_match = re.match(r'^(\s*)(\d+)[.)]\s+(.+)$', line)
        if number_match:
            indent = len(number_match.group(1))
            if indent >= 4:
                style = 'List Number 3'
            elif indent >= 2:
                style = 'List Number 2'
            else:
                style = 'List Number'
            p = doc.add_paragraph(style=style)
            parse_inline_formatting(p, number_match.group(3))
            idx += 1
            continue

        # --- Regular paragraph (merge consecutive non-special lines) ---
        para_lines = [stripped]
        idx += 1
        while idx < len(lines):
            next_stripped = lines[idx].strip()
            if not next_stripped or _is_special_line(next_stripped):
                break
            para_lines.append(next_stripped)
            idx += 1

        merged_text = ' '.join(para_lines)
        p = doc.add_paragraph()
        parse_inline_formatting(p, merged_text)

    # Save the document
    doc.save(output_path)
    print(f"Successfully converted: {input_path}")
    print(f"Output saved to: {output_path}")
    return output_path


def main():
    """Main entry point for command-line usage."""
    if len(sys.argv) < 2:
        print(__doc__)
        print("\nError: Please provide an input markdown file.")
        print("Usage: python md_to_docx.py input.md [output.docx]")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None

    if not os.path.exists(input_file):
        print(f"Error: Input file '{input_file}' not found.")
        sys.exit(1)

    try:
        convert_markdown_to_docx(input_file, output_file)
    except Exception as e:
        print(f"Error during conversion: {e}")
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
